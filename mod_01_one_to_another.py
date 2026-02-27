import streamlit as st
import pandas as pd
import json
import docx
from docx import Document
import PyPDF2
from io import BytesIO, StringIO
import base64
import zipfile

# Ορισμός τίτλου σελίδας
st.set_page_config(
    page_title="File Format Converter",
    page_icon="🔄",
    layout="wide"
)

# CSS για καλύτερη εμφάνιση
st.markdown("""
<style>
    .stDownloadButton button {
        width: 100%;
    }
    .file-info {
        padding: 10px;
        background-color: #f0f2f6;
        border-radius: 5px;
        margin: 5px 0;
    }
    .conversion-success {
        color: green;
        font-weight: bold;
        padding: 10px;
        background-color: #e8f5e8;
        border-radius: 5px;
    }
    .file-item {
        padding: 8px;
        margin: 4px 0;
        background-color: #f8f9fa;
        border-left: 4px solid #4e8cff;
        border-radius: 3px;
    }
</style>
""", unsafe_allow_html=True)

def read_docx(file):
    """Διαβάζει το περιεχόμενο από αρχείο DOCX"""
    doc = Document(file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def read_pdf(file):
    """Διαβάζει το περιεχόμενο από αρχείο PDF"""
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def create_docx_from_text(text, filename="converted.docx"):
    """Δημιουργεί αρχείο DOCX από κείμενο"""
    doc = Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    
    # Αποθήκευση σε buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_txt_from_text(text, filename="converted.txt"):
    """Δημιουργεί αρχείο TXT από κείμενο"""
    buffer = StringIO()
    buffer.write(text)
    buffer.seek(0)
    return BytesIO(buffer.getvalue().encode())

def process_file(file, conversion_type, options=None, custom_filename=None):
    """Επεξεργάζεται ένα αρχείο για μετατροπή"""
    if options is None:
        options = {}
    
    file_ext = file.name.split('.')[-1].lower()
    
    try:
        # Ανάγνωση αρχείου ανάλογα με τον τύπο
        if file_ext == 'docx':
            content = read_docx(file)
        elif file_ext == 'pdf':
            content = read_pdf(file)
        elif file_ext in ['txt']:
            content = file.read().decode('utf-8')
            file.seek(0)
        elif file_ext == 'json':
            json_content = json.load(file)
            content = json_content
            file.seek(0)
        elif file_ext == 'csv':
            file.seek(0)
            df = pd.read_csv(file)
            content = df
            file.seek(0)
        else:
            return None
        
        # Προκαθορισμένα ονόματα αρχείων ανά τύπο μετατροπής
        default_extensions = {
            "DOCX → TXT": ".txt",
            "TXT → DOCX": ".docx",
            "PDF → TXT": ".txt",
            "JSON → CSV": ".csv",
            "CSV → JSON": ".json",
            "TXT → PDF": ".pdf",
            "DOCX → PDF": ".pdf",
            "PDF → DOCX": ".docx"
        }
        
        # Μετατροπή
        output = None
        output_filename = file.name
        mime_type = "application/octet-stream"
        
        # DOCX → TXT
        if conversion_type == "DOCX → TXT" and file_ext == 'docx':
            output = create_txt_from_text(content)
            mime_type = "text/plain"
            
        # TXT → DOCX
        elif conversion_type == "TXT → DOCX" and file_ext == 'txt':
            output = create_docx_from_text(content)
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            
        # PDF → TXT
        elif conversion_type == "PDF → TXT" and file_ext == 'pdf':
            output = create_txt_from_text(content)
            mime_type = "text/plain"
            
        # JSON → CSV
        elif conversion_type == "JSON → CSV" and file_ext == 'json':
            df = pd.json_normalize(content)
            csv_buffer = StringIO()
            delimiter = options.get('delimiter', ',')
            encoding = options.get('encoding', 'utf-8')
            df.to_csv(csv_buffer, index=False, sep=delimiter)
            output = BytesIO(csv_buffer.getvalue().encode(encoding))
            mime_type = "text/csv"
            
        # CSV → JSON
        elif conversion_type == "CSV → JSON" and file_ext == 'csv':
            file.seek(0)
            delimiter = options.get('delimiter', ',')
            encoding = options.get('encoding', 'utf-8')
            df = pd.read_csv(file, delimiter=delimiter)
            json_str = df.to_json(orient='records', indent=2, force_ascii=False)
            output = BytesIO(json_str.encode(encoding))
            mime_type = "application/json"
            
        # TXT → PDF (placeholder - χρειάζεται επιπλέον βιβλιοθήκες)
        elif conversion_type == "TXT → PDF" and file_ext == 'txt':
            st.warning(f"Η μετατροπή TXT → PDF απαιτεί πρόσθετες βιβλιοθήκες. Το αρχείο '{file.name}' δεν μετατράπηκε.")
            return None
            
        # DOCX → PDF (placeholder)
        elif conversion_type == "DOCX → PDF" and file_ext == 'docx':
            st.warning(f"Η μετατροπή DOCX → PDF απαιτεί πρόσθετες βιβλιοθήκες. Το αρχείο '{file.name}' δεν μετατράπηκε.")
            return None
            
        # PDF → DOCX (placeholder)
        elif conversion_type == "PDF → DOCX" and file_ext == 'pdf':
            st.warning(f"Η μετατροπή PDF → DOCX απαιτεί πρόσθετες βιβλιοθήκες. Το αρχείο '{file.name}' δεν μετατράπηκε.")
            return None
        
        # Εάν υπάρχει έξοδος, καθορίζουμε το όνομα αρχείου
        if output is not None:
            if custom_filename:
                # Χρήση προσαρμοσμένου ονόματος
                base_name = custom_filename
                if not base_name.endswith(default_extensions.get(conversion_type, '')):
                    base_name += default_extensions.get(conversion_type, '')
                output_filename = base_name
            else:
                # Προκαθορισμένο όνομα
                output_filename = output_filename.replace(f'.{file_ext}', default_extensions.get(conversion_type, ''))
        
        return {
            'original_name': file.name,
            'output_name': output_filename,
            'content': output,
            'mime_type': mime_type,
            'success': output is not None
        }
        
    except Exception as e:
        return {
            'original_name': file.name,
            'error': str(e),
            'success': False
        }

def create_zip_from_results(results, zip_name="converted_files.zip"):
    """Δημιουργεί ZIP αρχείο από τα αποτελέσματα"""
    zip_buffer = BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for result in results:
            if result['success']:
                zip_file.writestr(result['output_name'], result['content'].getvalue())
    
    zip_buffer.seek(0)
    return zip_buffer

def main_one_to_another():
    
    # Δημιουργία δύο columns όπως ζητήσατε
    col1, col2 = st.columns(([4,1]), gap="small", vertical_alignment="top", border=True)
    
    with col2:
        st.subheader("Control Panel")
        st.markdown("---")
        
        # Επιλογή τύπου μετατροπής
        conversion_type = st.selectbox(
            "Τύπος Μετατροπής",
            [
                "DOCX → TXT",
                "TXT → DOCX",
                "PDF → TXT", 
                "TXT → PDF",
                "JSON → CSV",
                "CSV → JSON",
                "DOCX → PDF",
                "PDF → DOCX"
            ]
        )
        
        # Διαχωριστικό
        st.markdown("---")
        
        # Upload πολλαπλών αρχείων
        uploaded_files = st.file_uploader(
            "Upload file/s",
            type=['docx', 'txt', 'pdf', 'json', 'csv'],
            accept_multiple_files=True,
            help="Upload one or more files"
        )
        
        # Επιλογές για εξαγωγή
        st.markdown("---")
        st.markdown("### Extract Options")
        
        # Προσθήκη επιλογών ανάλογα με τον τύπο μετατροπής
        if "TXT" in conversion_type and ("DOCX" in conversion_type or "PDF" in conversion_type):
            add_page_numbers = st.checkbox("Add number of pages", value=False)
            font_size = st.slider("Font size", 8, 20, 12)
        
        if "JSON" in conversion_type or "CSV" in conversion_type:
            delimiter = st.selectbox("Seperator CSV", [",", ";", "\t"], index=0)
            encoding = st.selectbox("Coding", ["utf-8", "iso-8859-7", "windows-1253"], index=0)
        
        st.markdown("---")
        
        # ΠΕΡΙΟΧΗ ΓΙΑ ΚΑΘΟΡΙΣΜΟ ΟΝΟΜΑΤΩΝ ΑΡΧΕΙΩΝ
        st.markdown("### File Naming")
        
        # Προσθήκη δυνατότητας προσαρμογής ονόματος
        use_custom_name = st.checkbox("Use custom file name", value=False)
        
        custom_filename = None
        zip_custom_name = None
        
        if use_custom_name:
            if uploaded_files and len(uploaded_files) == 1:
                # Για ένα αρχείο: εισαγωγή ονόματος
                default_name = uploaded_files[0].name.split('.')[0]
                custom_filename = st.text_input(
                    "Output file name (without extension)",
                    value=default_name,
                    help="Enter the desired name for the converted file"
                )
            elif uploaded_files and len(uploaded_files) > 1:
                # Για πολλά αρχεία: εισαγωγή ονόματος ZIP
                zip_custom_name = st.text_input(
                    "ZIP file name (without .zip)",
                    value="converted_files",
                    help="Enter the desired name for the ZIP file"
                )
        
        st.markdown("---")
        
        # Κουμπί μετατροπής
        convert_button = st.button("Convert", use_container_width=True, type="primary")
        
        # Κουμπί καθαρισμού
        if st.button("Clean", use_container_width=True):
            st.rerun()
    
    with col1:
        st.subheader("Show Results")
        st.markdown("---")
        
        if uploaded_files:
            # Εμφάνιση λίστας αρχείων
            st.markdown(f"##### Uploaded Files ({len(uploaded_files)})")
            
            for i, file in enumerate(uploaded_files):
                file_details = {
                    "Όνομα": file.name,
                    "Τύπος": file.type,
                    "Μέγεθος": f"{file.size / 1024:.2f} KB"
                }
                
                st.markdown(f"""
                <div class="file-item">
                <b>{i+1}. {file.name}</b><br>
                <small>Τύπος: {file.type} | Μέγεθος: {file.size / 1024:.2f} KB</small>
                </div>
                """, unsafe_allow_html=True)
            
            # Προβολή περιεχομένου του πρώτου αρχείου
            if len(uploaded_files) == 1:
                file = uploaded_files[0]
                st.markdown("---")
                st.markdown("##### File content")
                
                try:
                    content = None
                    file_ext = file.name.split('.')[-1].lower()
                    
                    if file_ext == 'docx':
                        content = read_docx(file)
                        st.text_area("DOCX Content", content[:5000], height=300, key="docx_preview")
                        
                    elif file_ext == 'pdf':
                        content = read_pdf(file)
                        st.text_area("PDF content", content[:5000], height=300, key="pdf_preview")
                        
                    elif file_ext == 'txt':
                        content = file.read().decode()
                        file.seek(0)
                        st.text_area("TXT content", content[:5000], height=300, key="txt_preview")
                        
                    elif file_ext == 'json':
                        json_content = json.load(file)
                        file.seek(0)
                        st.json(json_content, expanded=False)
                        
                    elif file_ext == 'csv':
                        file.seek(0)
                        df = pd.read_csv(file)
                        st.dataframe(df.head(20), use_container_width=True)
                    
                except Exception as e:
                    st.error(f"❌ Error reading file: {str(e)}")
            
            # Μετατροπή όταν πατηθεί το κουμπί
            if convert_button:
                st.markdown("---")
                st.markdown("#### ⚙️ Conversion...")
                
                results = []
                options = {
                    'delimiter': delimiter if 'delimiter' in locals() else ',',
                    'encoding': encoding if 'encoding' in locals() else 'utf-8'
                }
                
                progress_bar = st.progress(0)
                
                # Επεξεργασία όλων των αρχείων
                for i, file in enumerate(uploaded_files):
                    progress_bar.progress((i + 1) / len(uploaded_files))
                    
                    # Χρήση προσαρμοσμένου ονόματος μόνο για πρώτο αρχείο
                    current_custom_name = custom_filename if (i == 0 and custom_filename) else None
                    
                    result = process_file(file, conversion_type, options, current_custom_name)
                    if result:
                        results.append(result)
                
                # Εμφάνιση αποτελεσμάτων
                st.markdown("---")
                st.markdown("##### Conversion results")
                
                successful = [r for r in results if r['success']]
                failed = [r for r in results if not r['success']]
                
                if successful:
                    st.success(f"✅ Successful conversions: {len(successful)}")
                    
                    # Λίστα επιτυχημένων
                    for result in successful:
                        st.markdown(f"✓ **{result['original_name']}** → **{result['output_name']}**")
                    
                    # Μετατροπή πολλαπλών αρχείων
                    if len(successful) > 1:
                        st.markdown("---")
                        st.markdown("##### File Package")
                        
                        # Δημιουργία ZIP με προσαρμοσμένο όνομα
                        zip_name = f"{zip_custom_name}.zip" if zip_custom_name else f"converted_files_{len(successful)}.zip"
                        zip_buffer = create_zip_from_results(successful, zip_name)
                        
                        # Προσαρμοσμένο κουμπί λήψης ZIP
                        download_label = f"📥 Download all as '{zip_name}'"
                        if not zip_custom_name:
                            download_label = f"📥 Download all ({len(successful)} files)"
                            
                        st.download_button(
                            label=download_label,
                            data=zip_buffer,
                            file_name=zip_name,
                            mime="application/zip",
                            use_container_width=True
                        )
                        
                        # Μεμονωμένα αρχεία
                        st.markdown("---")
                        st.markdown("##### Individual Files")
                        
                        for result in successful:
                            col_dl1, col_dl2 = st.columns([3, 1])
                            with col_dl1:
                                st.markdown(f"**{result['output_name']}**")
                            with col_dl2:
                                st.download_button(
                                    label="Download",
                                    data=result['content'].getvalue(),
                                    file_name=result['output_name'],
                                    mime=result['mime_type'],
                                    key=f"dl_{result['output_name']}_{hash(result['output_name'])}",
                                    use_container_width=True
                                )
                    
                    # Μετατροπή ενός αρχείου
                    elif len(successful) == 1:
                        result = successful[0]
                        
                        # Προβολή περιεχομένου για ορισμένες μετατροπές
                        if conversion_type in ["DOCX → TXT", "PDF → TXT"]:
                            content = result['content'].getvalue().decode('utf-8')
                            st.text_area("Content after conversion", content[:5000], height=300)
                        
                        elif conversion_type == "JSON → CSV":
                            csv_content = result['content'].getvalue().decode('utf-8')
                            df = pd.read_csv(StringIO(csv_content))
                            st.dataframe(df.head(20), use_container_width=True)
                        
                        elif conversion_type == "CSV → JSON":
                            json_content = json.loads(result['content'].getvalue().decode('utf-8'))
                            st.json(json_content, expanded=False)
                        
                        # Κουμπί λήψης με προσαρμοσμένο όνομα
                        st.markdown("---")
                        st.markdown("##### 📥 Download file")
                        
                        # Εμφάνιση ονόματος αρχείου
                        st.info(f"File will be saved as: **{result['output_name']}**")
                        
                        st.download_button(
                            label=f"Download '{result['output_name']}'",
                            data=result['content'].getvalue(),
                            file_name=result['output_name'],
                            mime=result['mime_type'],
                            use_container_width=True
                        )
                
                if failed:
                    st.markdown("---")
                    st.error(f"❌ Failed conversions: {len(failed)}")
                    
                    with st.expander("View errors"):
                        for result in failed:
                            st.markdown(f"**{result['original_name']}**: {result.get('error', 'Unknown error')}")
        
        else:
            # Οδηγίες όταν δεν υπάρχει αρχείο
            st.info("ℹ️ Please upload one or more files from the Control Panel to get started.")
            
            st.markdown("---")
            st.markdown("### 💡 Custom File Naming Tips")
            st.markdown("""
            1. Check **'Use custom file name'** in Control Panel
            2. For single files: Enter your desired name (without extension)
            3. For multiple files: Enter ZIP file name
            4. The extension will be added automatically based on conversion type
            5. If not specified, original names will be used with new extensions
            """)

